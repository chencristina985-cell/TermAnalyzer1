import os
import sys
import re
import sqlite3
import zipfile
import time
import threading
import queue
import collections
import xml.etree.ElementTree as ET
import json

# Tkinter UI 相关库
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

# 尝试导入第三方库以支持 docx 和 pdf 提取，失败时会有优雅提示
try:
    import docx
    from docx.shared import RGBColor, Pt, Cm
    from docx.enum.text import WD_COLOR_INDEX
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
except ImportError:
    docx = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

# 尝试导入 Win32 COM 库以支持一键打开原文自动搜索定位
try:
    import win32com.client
    import pythoncom
    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False

# 常量定义
if getattr(sys, 'frozen', False):
    EXE_DIR = os.path.dirname(sys.executable)
else:
    EXE_DIR = os.path.dirname(os.path.abspath(__file__))

DB_NAME = os.path.join(EXE_DIR, "term_analyzer.db")
DEFAULT_FONT = ("Microsoft YaHei", 10)
BOLD_FONT = ("Microsoft YaHei", 10, "bold")
TITLE_FONT = ("Microsoft YaHei", 11, "bold")
HEADER_FONT = ("Microsoft YaHei", 16, "bold")

# 定义性句子评分关键词及其权重
DEF_KEYWORDS = {
    "是指": 10,
    "是指的": 10,
    "定义为": 8,
    "定义": 6,
    "所谓": 8,
    "意味着": 6,
    "就是": 5,
    "是": 3,
    "关键": 5,
    "最重要": 5
}

# 常见停用词，用于常用搭配过滤
STOPWORDS = set("的了在是有和与及或一个于以之而等及着但也这那由去后前中里外上下一二三四五六七八九十")

# 句分割正则，支持中英文常见标点，并智能向后包含闭合引号、括号以保持句尾完整
SENTENCE_SPLIT_RE = re.compile(r'([。！？；\n!?;\r][”』）"\'\s]*)')

# ==========================================
# 1. 文本提取模块
# ==========================================

def extract_docx(filepath):
    """
    从 Word (.docx) 提取段落文本，记录大纲标题。
    篇目标题精准识别引擎（结合底层 XML 大纲级别 w:outlineLvl、居中对齐 w:jc、样式名及卷目特定规则）：
      - 卷一：2层级且“第...篇”样式居中 => “1层级标题·2层级标题”；如“基督在众教会...”等非“第...篇”样式，向下提取 3层级居中标题 => “3层级标题”
      - 卷二：3层级居中标题（自动排除“目录”） => 篇目标题
      - 卷三：2层级居中标题 => 篇目标题
      - 卷四：必须是 1层级居中标题 => 篇目标题（严格排除2层级）
      - 卷五：必须是包含“带领工人的职责”的标题 => 篇目标题
      - 卷六/卷七：除了居中外，必须包含“什么是/怎样/为什么要追求真理” => 篇目标题
      - 卷八：1层级居中标题 => 篇目标题
      - 其余文档：沿用最近的 2/3 层级标题
    """
    paragraphs_data = []
    try:
        doc = docx.Document(filepath)
        fname = os.path.basename(filepath)

        is_juan_1  = "卷一" in fname
        is_juan_2  = "卷二" in fname
        is_juan_3  = "卷三" in fname
        is_juan_4  = "卷四" in fname
        is_juan_5  = "卷五" in fname
        is_juan_67 = ("卷六" in fname) or ("卷七" in fname)
        is_juan_8  = "卷八" in fname

        # 第一组：卷一/二/三/六/七/八 → 小三号 = 15pt；第二组：卷四/五 → 三号 = 16pt
        if is_juan_4 or is_juan_5:
            TARGET_PT = 16.0
        else:
            TARGET_PT = 15.0
        PT_TOLERANCE = 0.6  # 允许±0.6pt差异容错

        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        def _get_font_size_pt(para):
            """\u83b7取段落实际字号(pt)\uff1a优先直接读取 run 字号\uff0c其次从底层 XML w:sz 读取\uff0c最后检查段落样式"""
            # 1. 优先从 run 读取
            for run in para.runs:
                if run.font.size:
                    return run.font.size.pt
            # 2. 从底层 XML w:sz 读取（单位是半磅\uff0c除以2转换为 pt）
            try:
                sz_nodes = para._element.xpath('.//w:sz/@w:val')
                if sz_nodes:
                    return int(sz_nodes[0]) / 2.0
            except Exception:
                pass
            # 3. 从段落样式读取
            try:
                if para.style and para.style.font and para.style.font.size:
                    return para.style.font.size.pt
            except Exception:
                pass
            return None

        def _is_center(para):
            """\u5224断\u5c45\u4e2d\u5bf9\u9f50\uff1a\u7ed3\u5408 python-docx \u5c5e\u6027\u3001\u5e95\u5c42 XML w:jc \u53ca\u6837\u5f0f\u540d\u79f0"""
            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                return True
            try:
                jc_vals = para._element.xpath('.//w:jc/@w:val')
                if jc_vals and "center" in [str(v).lower() for v in jc_vals]:
                    return True
            except Exception:
                pass
            try:
                if para.style and para.style.paragraph_format and para.style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    return True
            except Exception:
                pass
            if para.style and any(k in (para.style.name or "").lower() for k in ["居中", "center"]):
                return True
            return False

        latest_heading = ""
        current_nearest_heading = ""

        for para in doc.paragraphs:
            txt = para.text.strip()
            if not txt:
                continue

            font_pt = _get_font_size_pt(para)
            centered = _is_center(para)
            is_target_font = (font_pt is not None and abs(font_pt - TARGET_PT) <= PT_TOLERANCE)
            is_title_para = is_target_font and centered

            if is_juan_1:
                # 卷一：目标字号+居中 且包含“第...篇” → 最近大标题·本标题
                #          目标字号+居中 且不为“第...篇” → 仅用本标题
                if "基督在众教会行走时的说话" in txt and is_title_para:
                    continue
                if is_title_para:
                    is_di_pian = bool(re.search(r'第\s*[一二三四五六七八九十百千万0-9a-zA-Z]+\s*篇', txt))
                    if is_di_pian:
                        if latest_heading and latest_heading != txt:
                            current_nearest_heading = f"{latest_heading}·{txt}"
                        else:
                            current_nearest_heading = txt
                    else:
                        # 非“第...篇”居中标题：先记录为 latest_heading，等待更小字号的 3层内容标题
                        latest_heading = txt
                    continue
                # 小字号居中标题（三层级）：仅抓取本身
                if centered and font_pt is not None and font_pt < TARGET_PT - PT_TOLERANCE:
                    current_nearest_heading = txt
                    continue

            elif is_juan_2:
                # 卷二：目标字号+居中，不含“目”字
                if is_title_para and "目" not in txt:
                    current_nearest_heading = txt
                    continue
                elif is_title_para:
                    continue  # 含“目”字的直接跳过

            elif is_juan_3:
                # 卷三：目标字号+居中
                if is_title_para:
                    current_nearest_heading = txt
                    continue

            elif is_juan_4 or is_juan_5:
                # 卷四/卷五：目标字号+居中
                if is_title_para:
                    current_nearest_heading = txt
                    continue

            elif is_juan_67:
                # 卷六/卷七：目标字号+居中
                if is_title_para:
                    current_nearest_heading = txt
                    continue

            elif is_juan_8:
                # 卷八：目标字号+居中
                if is_title_para:
                    current_nearest_heading = txt
                    continue

            else:
                # 通用文档：居中标题就是篇目标题
                if is_title_para:
                    current_nearest_heading = txt
                    continue

            heading_label = current_nearest_heading if current_nearest_heading else "（无标题）"
            if len(txt) >= 10:
                paragraphs_data.append((txt, heading_label))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    ctxt = cell.text.strip()
                    if ctxt and len(ctxt) >= 10:
                        lbl = current_nearest_heading if current_nearest_heading else "（无标题）"
                        paragraphs_data.append((ctxt, lbl))
    except Exception as e:
        print(f"解析 DOCX 出错 {filepath}: {e}")
    return paragraphs_data


def extract_odt(filepath):
    """提取 OpenDocument (.odt) 文件的段落文本以及其所对应的大纲标题。
    采用字号+居中精确识别篇目标题：卷一/二/三/六/七/八=15pt，卷四/五=16pt。
    """
    paragraphs_data = []
    try:
        fname = os.path.basename(filepath)
        is_juan_1  = "卷一" in fname
        is_juan_2  = "卷二" in fname
        is_juan_3  = "卷三" in fname
        is_juan_4  = "卷四" in fname
        is_juan_5  = "卷五" in fname
        is_juan_67 = ("卷六" in fname) or ("卷七" in fname)
        is_juan_8  = "卷八" in fname

        if is_juan_4 or is_juan_5:
            TARGET_PT = 16.0
        else:
            TARGET_PT = 15.0
        PT_TOLERANCE = 0.6

        with zipfile.ZipFile(filepath) as z:
            content_xml = z.read('content.xml')

            ns_text  = '{urn:oasis:names:tc:opendocument:xmlns:text:1.0}'
            ns_style = '{urn:oasis:names:tc:opendocument:xmlns:style:1.0}'
            ns_fo    = '{urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0}'

            style_name_attr = f'{ns_text}style-name'

            # 解析所有样式的居中+字号信息
            style_info = {}  # name -> {'centered': bool, 'font_pt': float|None}

            def parse_styles(xml_bytes):
                try:
                    r = ET.fromstring(xml_bytes)
                    for style in r.iter(f'{ns_style}style'):
                        sname = style.attrib.get(f'{ns_style}name')
                        if not sname:
                            continue
                        info = style_info.get(sname, {'centered': False, 'font_pt': None})
                        pp = style.find(f'{ns_style}paragraph-properties')
                        if pp is not None:
                            align = pp.attrib.get(f'{ns_fo}text-align')
                            if align == 'center':
                                info['centered'] = True
                        tp = style.find(f'{ns_style}text-properties')
                        if tp is not None:
                            fs = tp.attrib.get(f'{ns_fo}font-size')
                            if fs:
                                fs = fs.strip()
                                if fs.endswith('pt'):
                                    try:
                                        info['font_pt'] = float(fs[:-2])
                                    except ValueError:
                                        pass
                        style_info[sname] = info
                except Exception:
                    pass

            parse_styles(content_xml)
            if 'styles.xml' in z.namelist():
                parse_styles(z.read('styles.xml'))

            root = ET.fromstring(content_xml)
            latest_heading = ""
            current_nearest_heading = ""

            for elem in root.iter():
                tag = elem.tag
                if not (tag.endswith('}p') or tag.endswith('}h')):
                    continue
                txt = "".join(elem.itertext()).strip()
                if not txt:
                    continue

                sname = elem.attrib.get(style_name_attr, "")
                info = style_info.get(sname, {})
                centered = info.get('centered', False)
                font_pt  = info.get('font_pt', None)
                if any(k in sname.lower() for k in ["居中", "center"]):
                    centered = True

                is_target_font = (font_pt is not None and abs(font_pt - TARGET_PT) <= PT_TOLERANCE)
                is_title_para  = is_target_font and centered

                if is_juan_1:
                    if "基督在众教会行走时的说话" in txt and is_title_para:
                        continue
                    if is_title_para:
                        is_di_pian = bool(re.search(r'第\s*[一二三四五六七八九十百千万0-9a-zA-Z]+\s*篇', txt))
                        if is_di_pian:
                            if latest_heading and latest_heading != txt:
                                current_nearest_heading = f"{latest_heading}·{txt}"
                            else:
                                current_nearest_heading = txt
                        else:
                            latest_heading = txt
                        continue
                    if centered and font_pt is not None and font_pt < TARGET_PT - PT_TOLERANCE:
                        current_nearest_heading = txt
                        continue

                elif is_juan_2:
                    if is_title_para and "目" not in txt:
                        current_nearest_heading = txt
                        continue
                    elif is_title_para:
                        continue

                elif is_juan_3:
                    if is_title_para:
                        current_nearest_heading = txt
                        continue

                elif is_juan_4 or is_juan_5:
                    if is_title_para:
                        current_nearest_heading = txt
                        continue

                elif is_juan_67:
                    if is_title_para:
                        current_nearest_heading = txt
                        continue

                elif is_juan_8:
                    if is_title_para:
                        current_nearest_heading = txt
                        continue

                else:
                    if is_title_para:
                        current_nearest_heading = txt
                        continue

                heading_label = current_nearest_heading if current_nearest_heading else "（无标题）"
                if len(txt) >= 10:
                    paragraphs_data.append((txt, heading_label))

    except Exception as e:
        print(f"解析 odt 出错 {filepath}: {e}")
    return paragraphs_data

def extract_pdf(filepath):
    """提取 PDF 文件的文本并切分成段落（因PDF无大纲标记，统一归为PDF无层级），自动跳过不可复制的扫描版文件"""
    if PdfReader is None:
        raise ImportError("未安装 pypdf 模块，无法解析 PDF 文件！请运行 pip install pypdf")
    paragraphs = []
    try:
        reader = PdfReader(filepath)
        # 性能优化：检查前两页是否有任何可复制文本，如果均无文本则判定为扫描版或受保护的不可复制 PDF，直接跳过以节省时间
        test_text = ""
        for page_idx in range(min(2, len(reader.pages))):
            page_text = reader.pages[page_idx].extract_text()
            if page_text:
                test_text += page_text.strip()
        if not test_text:
            print(f"提示: 检测到扫描版/不可复制的 PDF，已跳过解析: {filepath}")
            return []
            
        for page in reader.pages:
            text = page.extract_text()
            if text:
                # 根据换行符初步切分段落，或合并由于换行导致的硬折行
                lines = text.split('\n')
                current_p = []
                for line in lines:
                    line = line.strip()
                    if not line:
                        if current_p:
                            paragraphs.append(" ".join(current_p))
                            current_p = []
                        else:
                            pass
                    else:
                        current_p.append(line)
                if current_p:
                    paragraphs.append(" ".join(current_p))
    except Exception as e:
        print(f"解析 pdf 失败 {filepath}: {e}")
    # 统一附带 fallback 大纲名称
    return [(p, "（PDF无大纲层级）") for p in paragraphs]

def extract_text_from_file(filepath):
    """根据文件后缀调用对应的标题级提取器"""
    _, ext = os.path.splitext(filepath.lower())
    if ext == '.docx':
        return extract_docx(filepath)
    elif ext == '.odt':
        return extract_odt(filepath)
    elif ext == '.pdf':
        return extract_pdf(filepath)
    else:
        return []

# ==========================================
# 2. 数据库与索引模块
# ==========================================

class IndexDatabase:
    """SQLite 索引管理数据库，包含文件元数据和段落数据"""
    def __init__(self, db_path=DB_NAME):
        self.db_path = db_path
        self.conn = None
        self._init_db()

    def _init_db(self):
        self.conn = sqlite3.connect(self.db_path, check_same_thread=False)
        cursor = self.conn.cursor()
        
        # SQLite 性能调优，大幅降低低配电脑上的读写卡顿
        try:
            cursor.execute("PRAGMA journal_mode=WAL")
            cursor.execute("PRAGMA synchronous=NORMAL")
            cursor.execute("PRAGMA cache_size=-50000") # 约 50MB 内存缓存
        except Exception:
            pass
            
        # 1. 创建文件表（存路径、修改时间、以及所属子文件夹分类，实现增量更新与分类筛选）
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS files (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                filepath TEXT UNIQUE,
                filename TEXT,
                file_type TEXT,
                mtime REAL,
                category TEXT
            )
        """)
        
        # 自动数据库结构迁移：尝试为旧数据库添加 category 字段以确保兼容性
        try:
            cursor.execute("ALTER TABLE files ADD COLUMN category TEXT")
        except sqlite3.OperationalError:
            pass  # 字段已存在
            
        # 2. 创建段落数据表 (增添 heading 字段，保存所属的大纲级别为 2 的标题)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS paragraphs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                file_id INTEGER,
                paragraph_index INTEGER,
                content TEXT,
                heading TEXT,
                FOREIGN KEY(file_id) REFERENCES files(id) ON DELETE CASCADE
            )
        """)
        
        # 自动数据库结构迁移：尝试为段落表添加 heading 字段以确保向后兼容性
        try:
            cursor.execute("ALTER TABLE paragraphs ADD COLUMN heading TEXT")
        except sqlite3.OperationalError:
            pass  # 字段已存在
        
        # 3. 创建索引加速
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_file_id ON paragraphs(file_id)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_content ON paragraphs(content)")
        
        self.conn.commit()

    def get_indexed_files(self):
        """获取所有已索引的文件路径、修改时间及分类"""
        cursor = self.conn.cursor()
        cursor.execute("SELECT filepath, mtime, category FROM files")
        return {row[0]: (row[1], row[2]) for row in cursor.fetchall()}

    def update_file_category(self, filepath, category):
        """仅更新文件所属的分类信息"""
        cursor = self.conn.cursor()
        try:
            cursor.execute("UPDATE files SET category = ? WHERE filepath = ?", (category, filepath))
            self.conn.commit()
        except Exception as e:
            print(f"更新文件分类失败: {e}")

    def add_file(self, filepath, paragraphs, category="根目录"):
        """插入或更新文件，并写入段落数据和所属子文件夹分类"""
        cursor = self.conn.cursor()
        filename = os.path.basename(filepath)
        _, ext = os.path.splitext(filepath.lower())
        file_type = ext.replace('.', '')
        mtime = os.path.getmtime(filepath)
        
        try:
            # 事务开始
            cursor.execute("BEGIN TRANSACTION")
            
            # 删除旧数据
            cursor.execute("SELECT id FROM files WHERE filepath = ?", (filepath,))
            row = cursor.fetchone()
            if row:
                file_id = row[0]
                cursor.execute("DELETE FROM paragraphs WHERE file_id = ?", (file_id,))
                cursor.execute("UPDATE files SET mtime = ?, category = ? WHERE id = ?", (mtime, category, file_id))
            else:
                cursor.execute("""
                    INSERT INTO files (filepath, filename, file_type, mtime, category) 
                    VALUES (?, ?, ?, ?, ?)
                """, (filepath, filename, file_type, mtime, category))
                file_id = cursor.lastrowid
                
            # 批量插入段落 (paragraphs 为含有段落文本和归属大纲2标题的 tuple 列表)
            insert_data = [
                (file_id, idx, text, heading)
                for idx, (text, heading) in enumerate(paragraphs)
            ]
            cursor.executemany("""
                INSERT INTO paragraphs (file_id, paragraph_index, content, heading)
                VALUES (?, ?, ?, ?)
            """, insert_data)
            
            self.conn.commit()
        except Exception as e:
            self.conn.rollback()
            raise e

    def remove_file(self, filepath):
        """当本地文件被删除时，从数据库中同步清理"""
        cursor = self.conn.cursor()
        try:
            cursor.execute("DELETE FROM files WHERE filepath = ?", (filepath,))
            self.conn.commit()
        except Exception as e:
            print(f"清理失效文件记录失败 {filepath}: {e}")

    def get_stats(self):
        """获取索引库的统计信息"""
        cursor = self.conn.cursor()
        cursor.execute("SELECT COUNT(*) FROM files")
        file_count = cursor.fetchone()[0]
        cursor.execute("SELECT COUNT(*) FROM paragraphs")
        para_count = cursor.fetchone()[0]
        return file_count, para_count

    def get_all_categories(self):
        """获取数据库中目前已存在的所有文件夹分类列表"""
        cursor = self.conn.cursor()
        cursor.execute("SELECT DISTINCT category FROM files WHERE category IS NOT NULL")
        cats = [row[0] for row in cursor.fetchall() if row[0]]
        # 排序，让“根目录”排在前面，其余按字典排序
        if "根目录" in cats:
            cats.remove("根目录")
            return ["根目录"] + sorted(cats)
        return sorted(cats)

    def search_paragraphs(self, term, categories=None):
        """支持空格/斜杠/竖线分隔的多同义词 OR 匹配检索"""
        cursor = self.conn.cursor()
        terms = [t.strip() for t in re.split(r"[|/,\s，]+", term) if t.strip()]
        if not terms:
            terms = [term]
        like_sql = " OR ".join(["p.content LIKE ?"] * len(terms))
        params = [f"%{t}%" for t in terms]
        if categories:
            ph = ",".join("?" for _ in categories)
            sql = f"""
                SELECT p.content, f.filepath, f.filename, p.paragraph_index, p.heading
                FROM paragraphs p JOIN files f ON p.file_id = f.id
                WHERE ({like_sql}) AND f.category IN ({ph})
            """
            cursor.execute(sql, params + list(categories))
        else:
            sql = f"""
                SELECT p.content, f.filepath, f.filename, p.paragraph_index, p.heading
                FROM paragraphs p JOIN files f ON p.file_id = f.id
                WHERE {like_sql}
            """
            cursor.execute(sql, params)
        return cursor.fetchall()

    def close(self):
        if self.conn:
            self.conn.close()

# ==========================================
# 3. 核心算法：术语定义评分与搭配分析
# ==========================================

def split_into_sentences(text):
    """将段落拆分为句子，并保持标点符号的完整性"""
    tokens = SENTENCE_SPLIT_RE.split(text)
    sentences = []
    # 重新拼接标点符号
    for i in range(0, len(tokens) - 1, 2):
        sentence = tokens[i] + tokens[i+1]
        sentence = sentence.strip()
        if sentence:
            sentences.append(sentence)
    if len(tokens) % 2 == 1:
        last = tokens[-1].strip()
        if last:
            sentences.append(last)
    return sentences

def score_query_in_paragraph(sentences, i, query):
    """对第 i 句中出现的 query，结合紧邻的特征词（距离为0）进行加权评分"""
    current_sent = sentences[i]
    score = 0.0
    matches_info = []
    has_close_definition = False
    q_len = len(query)
    best_close_kw_bonus = 0.0
    best_kw_info = ""
    start_idx = 0
    while True:
        pos = current_sent.find(query, start_idx)
        if pos == -1:
            break
        q_end = pos + q_len
        has_suowei = (pos - 2 >= 0 and current_sent.startswith("所谓", pos - 2))
        has_jiushi = current_sent.startswith("就是", q_end)
        if has_suowei and has_jiushi:
            best_close_kw_bonus = 50.0
            best_kw_info = "所谓+术语+就是 (第一权重) (+50.0)"
            has_close_definition = True
            start_idx = pos + 1
            continue
        type_a = {"是指":15.,"就是":12.,"意味着":12.,"关键是":12.,"最重要":10.,"是":8.,"定义为":12.,"定性为":12.,"如何定义":15.,"如何定性":15.,"怎么定义":15.,"怎么定性":15.}
        for kw, w in type_a.items():
            if current_sent.startswith(kw, q_end):
                if w > best_close_kw_bonus:
                    best_close_kw_bonus = w
                    best_kw_info = f"术语+{kw} (紧邻) (+{w})"
                    has_close_definition = True
                break
        type_b = {"什么是":15.,"所谓":15.,"属于":10.,"就是":12.,"如何定义":15.,"如何定性":15.,"怎么定义":15.,"怎么定性":15.,"定义":10.,"定性":10.}
        for kw, w in type_b.items():
            kl = len(kw)
            cs = pos - kl
            if cs >= 0 and current_sent.startswith(kw, cs):
                if w > best_close_kw_bonus:
                    best_close_kw_bonus = w
                    best_kw_info = f"{kw}+术语 (紧邻) (+{w})"
                    has_close_definition = True
                break
        start_idx = pos + 1
    if has_close_definition:
        score += best_close_kw_bonus
        matches_info.append(best_kw_info)
        if i == 0:
            score += 5.0
            matches_info.append("首句定义 (+5)")
        elif i == 1:
            score += 2.0
            matches_info.append("次句定义 (+2)")
    return score, matches_info, has_close_definition


def analyze_and_rank_results(results, query, assoc2="", assoc3=""):
    """对段落进行评分和降序排序。支持同义词列表(query 拆分)以及第二、第三关联词优先加分。"""
    terms = [t.strip() for t in re.split(r"[|/,\s，]+", query) if t.strip()]
    if not terms:
        terms = [query]
    ranked_results = []
    for content, filepath, filename, para_idx, heading in results:
        sentences = split_into_sentences(content)
        total_def_score = 0.0
        def_count = 0
        best_sentence = ""
        best_info = []
        max_single_score = 0.0
        has_close_definition = False
        for idx, sent in enumerate(sentences):
            for t in terms:
                if t in sent:
                    sc, info, is_def = score_query_in_paragraph(sentences, idx, t)
                    if sc > max_single_score:
                        max_single_score = sc
                        best_sentence = sent
                        best_info = list(info)
                    if is_def:
                        total_def_score += sc
                        def_count += 1
                        has_close_definition = True
                    break
        if has_close_definition:
            max_score = total_def_score if def_count > 1 else max_single_score
            if def_count > 1:
                best_info.append(f"多重定义关联 (共 {def_count} 处)")
        else:
            max_score = 0.0
        tf = sum(content.count(t) for t in terms)
        char_count = max(150, len(content))
        tf_weight = round((tf / char_count) * 800.0 + tf * 2.0, 1)
        if not has_close_definition or max_score == 0.0:
            best_info = [f"术语频次: {tf}次 | 密度: {tf/max(1,len(content))*100:.1f}% (+{tf_weight})"]
            for sent in sentences:
                if any(t in sent for t in terms):
                    best_sentence = sent
                    break
            if not best_sentence:
                best_sentence = content[:30] + "..."
        else:
            best_info.insert(0, f"术语频次: {tf}次 | 密度: {tf/max(1,len(content))*100:.1f}% (+{tf_weight})")
        total_score = tf_weight + max_score
        if assoc2 and assoc2 in content:
            total_score += 100.0
            best_info.append(f"含第二关联词'{assoc2}' (+100)")
        if assoc3 and assoc3 in content:
            total_score += 100.0
            best_info.append(f"含第三关联词'{assoc3}' (+100)")
        length_penalty = min(0.15, len(content) / 8000.0)
        final_score = total_score * (1.0 - length_penalty)
        ranked_results.append({
            "content": content, "filepath": filepath, "filename": filename,
            "para_index": para_idx, "heading": heading or "（无标题）",
            "score": final_score, "info": best_info,
            "matched_sentence": best_sentence, "has_close_definition": has_close_definition
        })
    ranked_results.sort(key=lambda x: x["score"], reverse=True)
    return ranked_results


def calculate_collocations(paragraphs, query, window_size=10):
    """计算术语（含同义词）周围常见搭配词 (2-4字)。"""
    colloc_counter = collections.Counter()
    punct = re.compile(r'[，。！？、；：“”‘’（）()[\\]{}<>《》\\s\\-\\_.,;:!?()]' )
    terms = [t.strip() for t in re.split(r"[|/,\s，]+", query) if t.strip()]
    if not terms:
        terms = [query]
    for text_para in paragraphs:
        for q in terms:
            start = 0
            while True:
                idx = text_para.find(q, start)
                if idx == -1:
                    break
                left  = text_para[max(0, idx-window_size):idx]
                right = text_para[idx+len(q):min(len(text_para), idx+len(q)+window_size)]
                for ctx in [left, right]:
                    for seg in punct.split(ctx):
                        seg = seg.strip()
                        if not seg:
                            continue
                        for n in [2,3,4]:
                            if len(seg) >= n:
                                for i in range(len(seg)-n+1):
                                    ng = seg[i:i+n]
                                    if all(c in STOPWORDS for c in ng):
                                        continue
                                    if ng[0] in STOPWORDS and len(ng)==2:
                                        continue
                                    colloc_counter[ng] += 1
                start = idx + len(q)
    return colloc_counter.most_common(15)



def highlight_term_and_close_keywords(text_widget, query, query_tag, def_kw_tag):
    """
    在 Tkinter Text 控件中高亮检索术语（黄底）及其紧邻的定义特征词（红色加粗）。
    支持多同义词（空格/斜杠/竖线分隔）的 OR 匹配。
    """
    # 拆分同义词
    terms = [t.strip() for t in re.split(r'[|/,\s，]+', query) if t.strip()] or [query]

    # 配置 tag 样式
    text_widget.tag_configure(query_tag, background="#fef08a", font=("Microsoft YaHei", 10, "bold"))
    text_widget.tag_configure(def_kw_tag, foreground="#dc2626", font=("Microsoft YaHei", 10, "bold"))

    def_kws = ["是指", "就是", "意味着", "关键是", "最重要", "是", "定义为", "定性为",
               "什么是", "所谓", "属于", "定义", "定性", "如何定义", "如何定性", "怎么定义", "怎么定性"]

    for q in terms:
        if not q:
            continue
        start = "1.0"
        while True:
            pos = text_widget.search(q, start, stopindex="end", nocase=False)
            if not pos:
                break
            end_pos = f"{pos}+{len(q)}c"
            text_widget.tag_add(query_tag, pos, end_pos)

            # 检查右侧紧邻定义词
            for kw in def_kws:
                kw_pos = end_pos
                kw_end = f"{kw_pos}+{len(kw)}c"
                segment = text_widget.get(kw_pos, kw_end)
                if segment == kw:
                    text_widget.tag_add(def_kw_tag, kw_pos, kw_end)
                    break

            # 检查左侧紧邻定义词
            for kw in def_kws:
                kw_end_pos = pos
                kw_start_pos = f"{pos}-{len(kw)}c"
                try:
                    segment = text_widget.get(kw_start_pos, kw_end_pos)
                    if segment == kw:
                        text_widget.tag_add(def_kw_tag, kw_start_pos, kw_end_pos)
                        break
                except Exception:
                    pass

            start = end_pos


class ModernCard(tk.Frame):

    """自定义卡片小组件，用于美观展示段落结果"""
    def __init__(self, parent, result, query, select_callback, open_callback, **kwargs):
        super().__init__(parent, bg="#ffffff", bd=1, relief="solid", highlightthickness=0, **kwargs)
        self.configure(highlightbackground="#e2e8f0")
        
        self.result = result
        self.filepath = result["filepath"]
        self.filename = result["filename"]
        
        # 鼠标悬停事件
        self.bind("<Enter>", self._on_enter)
        self.bind("<Leave>", self._on_leave)
        
        # 头部栏 (文件名 + 分数)
        header_frame = tk.Frame(self, bg="#f8fafc", padx=8, pady=4)
        header_frame.pack(fill="x", side="top")
        
        # 格式化并精简展示卡片的篇目大纲信息（清除小标题，仅保留主标题且不显示第几段）
        book_match = re.search(r'《([^》]+)》', self.filename)
        book_title = book_match.group(1) if book_match else os.path.splitext(self.filename)[0]
        
        clean_heading = result['heading'].split(" > ")[-1].strip()
        if clean_heading and clean_heading != "（无标题）":
            title_info = f"《{book_title}·{clean_heading}》"
        else:
            title_info = f"《{book_title}》"
            
        file_lbl = tk.Label(
            header_frame, 
            text=f"📄 {title_info}", 
            font=BOLD_FONT, 
            fg="#0ea5e9", 
            bg="#f8fafc", 
            cursor="hand2"
        )
        file_lbl.pack(side="left")
        file_lbl.bind("<Button-1>", lambda e: open_callback(self.filepath, result["content"]))
        
        # 评分勋章
        score_text = f"评分: {result['score']:.1f}"
        if result['score'] > 5:
            badge_bg = "#f0fdf4"  # 绿色背景
            badge_fg = "#15803d"
        else:
            badge_bg = "#f1f5f9"  # 灰色背景
            badge_fg = "#475569"
            
        score_badge = tk.Label(
            header_frame, 
            text=score_text, 
            font=("Microsoft YaHei", 9, "bold"), 
            bg=badge_bg, 
            fg=badge_fg, 
            padx=6, 
            pady=2
        )
        score_badge.pack(side="right")
        
        # 中部：分析加分原因
        if result["info"]:
            info_text = "加分原因: " + " | ".join(result["info"])
            info_lbl = tk.Label(
                self, 
                text=info_text, 
                font=("Microsoft YaHei", 8), 
                fg="#64748b", 
                bg="#ffffff", 
                anchor="w", 
                padx=10, 
                pady=2
            )
            info_lbl.pack(fill="x")
            
        # 下部：文本内容展示（高亮关键词）
        text_container = tk.Frame(self, bg="#ffffff", padx=8, pady=4)
        text_container.pack(fill="both", expand=True)
        
        self.text_widget = tk.Text(
            text_container, 
            wrap="word", 
            font=DEFAULT_FONT, 
            bg="#ffffff", 
            fg="#334155", 
            bd=0, 
            highlightthickness=0, 
            height=3, 
            padx=4, 
            pady=4,
            cursor="arrow"
        )
        self.text_widget.pack(fill="both", expand=True)
        
        # 生成展示切片 (Snippet)，确保搜索词在卡片中可见并被高亮
        snippet_text = self._generate_snippet(result["content"], result["matched_sentence"], query)
        self.text_widget.insert("1.0", snippet_text)
        self.text_widget.configure(state="disabled")
        
        # 对关键词添加标签样式并高亮 (仅高亮检索术语及紧邻 5 字以内的定义特征词)
        highlight_term_and_close_keywords(self.text_widget, query, "query", "def_kw")
        
        # 绑定点击卡片事件 (选中该卡片在右侧显示大图/详细分析)
        for w in (self, text_container, self.text_widget):
            w.bind("<Button-1>", lambda e: select_callback(self.result))
            
    def _on_enter(self, event):
        self.configure(highlightbackground="#0ea5e9", bd=1)
        
    def _on_leave(self, event):
        self.configure(highlightbackground="#e2e8f0", bd=1)
                
    def _generate_snippet(self, content, matched_sentence, query):
        """从长段落中截取并拼合出匹配位置周围的上下文切片(Snippet)，支持多同义词"""
        sentences = split_into_sentences(content)
        # 拆分为多个同义词词条
        terms = [t.strip() for t in re.split(r'[|/,\s，]+', query) if t.strip()] or [query]
        matched_idx = -1

        # 1. 优先寻找与评分时捕获的 matched_sentence 相同的句子
        for idx, sent in enumerate(sentences):
            if sent == matched_sentence:
                matched_idx = idx
                break

        # 2. 兜底：找包含任意一个同义词的句子
        if matched_idx == -1:
            for idx, sent in enumerate(sentences):
                if any(t in sent for t in terms):
                    matched_idx = idx
                    break

        # 3. 截取该句以及前一句、后一句作为 Snippet 展示
        if matched_idx != -1:
            start_idx = max(0, matched_idx - 1)
            end_idx = min(len(sentences), matched_idx + 2)
            snippet_sents = sentences[start_idx:end_idx]
            snippet_text = " ".join(snippet_sents)
            if start_idx > 0:
                snippet_text = "... " + snippet_text
            if end_idx < len(sentences):
                snippet_text = snippet_text + " ..."
            return snippet_text
        else:
            return content[:150] + "..." if len(content) > 150 else content

class ScrollableFrame(tk.Frame):
    """带滚动条的自适应 Frame 容器"""
    def __init__(self, container, *args, **kwargs):
        super().__init__(container, *args, **kwargs)
        self.canvas = tk.Canvas(self, borderwidth=0, highlightthickness=0, bg="#f8fafc")
        self.scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = tk.Frame(self.canvas, bg="#f8fafc")
        
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(
                scrollregion=self.canvas.bbox("all")
            )
        )
        
        self.canvas_window = self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        
        self.canvas.bind('<Configure>', self._on_canvas_configure)
        
        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")
        
        # 仅在鼠标停留在区域内时绑定滚轮事件
        self.canvas.bind("<Enter>", lambda e: self.canvas.bind_all("<MouseWheel>", self._on_mousewheel))
        self.canvas.bind("<Leave>", lambda e: self.canvas.unbind_all("<MouseWheel>"))
        
    def _on_canvas_configure(self, event):
        self.canvas.itemconfigure(self.canvas_window, width=event.width)
        
    def _on_mousewheel(self, event):
        self.canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
    def clear(self):
        for child in self.scrollable_frame.winfo_children():
            child.destroy()

class TermAnalyzerApp:
    """应用程序主窗口"""
    def __init__(self, root):
        self.root = root
        self.root.title("本地 Word 术语分析与检索工具")
        self.root.geometry("1180x750")
        self.root.minsize(1000, 600)
        
        self.db = IndexDatabase()
        self.indexing_queue = queue.Queue()
        self.is_indexing = False
        self.current_results = []
        
        # 加载持久化配置记录
        self._load_settings()
        
        # 设置样式
        self.style = ttk.Style()
        self.style.theme_use("clam")
        
        # 配置 ttk 样式使其更美观
        self.style.configure("TProgressbar", thickness=15, troughcolor="#e2e8f0", background="#0ea5e9")
        self.style.configure("Horizontal.TScrollbar", arrowsize=10, background="#cbd5e1")
        self.style.configure("Vertical.TScrollbar", arrowsize=10, background="#cbd5e1")
        
        # 主题色彩
        self.bg_color = "#f8fafc"
        self.root.configure(bg=self.bg_color)
        
        # 初始化 UI
        self._build_ui()
        
        # 如果有缓存的历史目录且该目录依然有效，直接加载并呈现分类
        if self.selected_folder and os.path.exists(self.selected_folder):
            self.dir_var.set(self.selected_folder)
            self._rebuild_category_ui()
            
        # 启动队列监听器
        self._listen_queue()
        
        # 更新状态栏
        self._update_db_stats()

    def _build_ui(self):
        # 1. 顶部控制面板（文件夹选择、索引状态、进度条）
        top_panel = tk.Frame(self.root, bg="#ffffff", bd=1, relief="solid", highlightthickness=0)
        top_panel.configure(highlightbackground="#e2e8f0")
        top_panel.pack(fill="x", padx=15, pady=10)
        
        # 文件夹选择区域
        folder_frame = tk.Frame(top_panel, bg="#ffffff")
        folder_frame.pack(fill="x", padx=15, pady=8)
        
        lbl_dir = tk.Label(folder_frame, text="目标文件夹:", font=BOLD_FONT, fg="#0f172a", bg="#ffffff")
        lbl_dir.pack(side="left", padx=2)
        
        self.dir_var = tk.StringVar(value="请选择一个本地文件夹并建立索引...")
        self.entry_dir = tk.Entry(
            folder_frame, 
            textvariable=self.dir_var, 
            font=DEFAULT_FONT, 
            width=50, 
            bg="#f1f5f9", 
            fg="#64748b",
            bd=1,
            relief="solid"
        )
        self.entry_dir.pack(side="left", padx=8, ipady=3, fill="x", expand=True)
        
        # 按钮：浏览
        btn_browse = tk.Button(
            folder_frame, 
            text="选择文件夹 📁", 
            command=self._select_folder, 
            font=DEFAULT_FONT, 
            bg="#f1f5f9", 
            fg="#0f172a", 
            relief="flat", 
            activebackground="#cbd5e1",
            cursor="hand2"
        )
        btn_browse.pack(side="left", padx=4)
        
        # 按钮：构建索引
        self.btn_index = tk.Button(
            folder_frame, 
            text="更新/扫描索引 ⚡", 
            command=self._start_indexing, 
            font=BOLD_FONT, 
            bg="#0f172a", 
            fg="#ffffff", 
            relief="flat",
            activebackground="#1e293b",
            activeforeground="#ffffff",
            cursor="hand2"
        )
        self.btn_index.pack(side="left", padx=4)
        
        # 进度状态与条形图
        self.progress_frame = tk.Frame(top_panel, bg="#ffffff")
        self.progress_frame.pack(fill="x", padx=15, pady=(0, 8))
        self.progress_frame.pack_forget()  # 默认隐藏
        
        self.lbl_progress_status = tk.Label(self.progress_frame, text="正在读取文件列表...", font=DEFAULT_FONT, bg="#ffffff", fg="#475569")
        self.lbl_progress_status.pack(side="top", anchor="w", pady=(0, 2))
        
        self.progress_bar = ttk.Progressbar(self.progress_frame, orient="horizontal", mode="determinate")
        self.progress_bar.pack(fill="x", side="left", expand=True)
        
        # 2. 搜索控制面板
        search_panel = tk.Frame(self.root, bg="#ffffff", bd=1, relief="solid", highlightthickness=0)
        search_panel.configure(highlightbackground="#e2e8f0")
        search_panel.pack(fill="x", padx=15, pady=(0, 10))
        
        search_frame = tk.Frame(search_panel, bg="#ffffff", padx=15, pady=10)
        search_frame.pack(fill="x")
        
        lbl_search = tk.Label(search_frame, text="搜索术语:", font=BOLD_FONT, fg="#0f172a", bg="#ffffff")
        lbl_search.pack(side="left", padx=2)
        
        self.search_var = tk.StringVar()
        self.entry_search = tk.Entry(
            search_frame, 
            textvariable=self.search_var, 
            font=("Microsoft YaHei", 11), 
            bg="#ffffff", 
            fg="#0f172a",
            bd=1,
            relief="solid",
            highlightthickness=1,
            highlightcolor="#0ea5e9"
        )
        self.entry_search.pack(side="left", padx=8, ipady=4, fill="x", expand=True)
        self.entry_search.bind("<Return>", lambda e: self._perform_search())
        
        # 按钮：搜索
        btn_search = tk.Button(
            search_frame, 
            text="开始检索 🔍", 
            command=self._perform_search, 
            font=BOLD_FONT, 
            bg="#0ea5e9", 
            fg="#ffffff", 
            relief="flat",
            activebackground="#0284c7",
            activeforeground="#ffffff",
            cursor="hand2",
            padx=12
        )
        btn_search.pack(side="left", padx=4)
        
        # 按钮：批量检索
        btn_batch_search = tk.Button(
            search_frame, 
            text="批量术语检索 📂", 
            command=self._batch_term_search_and_export, 
            font=BOLD_FONT, 
            bg="#10b981", 
            fg="#ffffff", 
            relief="flat",
            activebackground="#047857",
            activeforeground="#ffffff",
            cursor="hand2",
            padx=12
        )
        btn_batch_search.pack(side="left", padx=4)
        
        # 统计文本提示 (搜索用时等)
        self.lbl_search_stats = tk.Label(search_frame, text="已就绪", font=DEFAULT_FONT, fg="#64748b", bg="#ffffff")
        self.lbl_search_stats.pack(side="right", padx=10)
        
        # 第二行：第二/第三关联词优先搜索输入
        assoc_frame = tk.Frame(search_panel, bg="#ffffff", padx=15, pady=4)
        assoc_frame.pack(fill="x")
        
        lbl_assoc2 = tk.Label(assoc_frame, text="第二关联词 (优先):", font=DEFAULT_FONT, fg="#475569", bg="#ffffff")
        lbl_assoc2.pack(side="left", padx=2)
        
        self.assoc2_var = tk.StringVar()
        self.entry_assoc2 = tk.Entry(
            assoc_frame, 
            textvariable=self.assoc2_var, 
            font=DEFAULT_FONT, 
            bg="#ffffff", 
            fg="#0f172a",
            bd=1,
            relief="solid",
            width=18
        )
        self.entry_assoc2.pack(side="left", padx=8, ipady=2)
        
        lbl_assoc3 = tk.Label(assoc_frame, text="第三关联词 (优先):", font=DEFAULT_FONT, fg="#475569", bg="#ffffff")
        lbl_assoc3.pack(side="left", padx=20)
        
        self.assoc3_var = tk.StringVar()
        self.entry_assoc3 = tk.Entry(
            assoc_frame, 
            textvariable=self.assoc3_var, 
            font=DEFAULT_FONT, 
            bg="#ffffff", 
            fg="#0f172a",
            bd=1,
            relief="solid",
            width=18
        )
        self.entry_assoc3.pack(side="left", padx=8, ipady=2)
        
        lbl_hint = tk.Label(
            assoc_frame, 
            text="* 提示：搜索术语支持用空格/斜杠/竖线输入多个同义词以实现或(OR)匹配", 
            font=("Microsoft YaHei", 8), 
            fg="#ef4444", 
            bg="#ffffff"
        )
        lbl_hint.pack(side="right", padx=10)
        
        # 3. 主分割视窗 (左侧结果区，右侧分析区)
        main_split = tk.PanedWindow(self.root, orient="horizontal", bg="#e2e8f0", sashwidth=4, bd=0)
        main_split.pack(fill="both", expand=True, padx=15, pady=(0, 10))
        
        # 左侧列表容器
        left_container = tk.Frame(main_split, bg="#f8fafc")
        main_split.add(left_container, minsize=450, stretch="always")
        
        lbl_left_title = tk.Label(
            left_container, 
            text=" 🔍 检索结果（高权重定义优先）", 
            font=TITLE_FONT, 
            fg="#0f172a", 
            bg="#f8fafc", 
            anchor="w",
            pady=4
        )
        lbl_left_title.pack(fill="x")
        
        # 滚动结果框
        self.results_frame = ScrollableFrame(left_container)
        self.results_frame.pack(fill="both", expand=True)
        
        # 默认无结果占位提示
        self.lbl_placeholder = tk.Label(
            self.results_frame.scrollable_frame, 
            text="请输入搜索术语开始快速检索文档。\n搜索结果将智能提取定义，并高亮呈现高分段落。", 
            font=DEFAULT_FONT, 
            fg="#94a3b8", 
            bg="#f8fafc",
            pady=100
        )
        self.lbl_placeholder.pack(fill="both", expand=True)
        
        # 右侧侧边栏 (搭配词 + 段落详情)
        right_container = tk.Frame(main_split, bg="#ffffff", padx=10, pady=5)
        main_split.add(right_container, minsize=350, stretch="never")
        
        # --- 子文件夹分类筛选面板 ---
        category_title_lbl = tk.Label(right_container, text=" 📁 子文件夹分类筛选", font=TITLE_FONT, fg="#0f172a", bg="#ffffff", anchor="w")
        category_title_lbl.pack(fill="x", pady=(0, 5))
        
        self.category_frame = ScrollableFrame(right_container)
        self.category_frame.canvas.configure(height=140)
        self.category_frame.pack(fill="x", pady=(0, 15))
        self.category_vars = {}
        
        # --- 段落详细内容面板 ---
        detail_title_lbl = tk.Label(right_container, text=" 📑 选中段落详细信息", font=TITLE_FONT, fg="#0f172a", bg="#ffffff", anchor="w")
        detail_title_lbl.pack(fill="x", pady=(0, 5))
        
        detail_card = tk.Frame(right_container, bg="#ffffff", bd=1, relief="solid", highlightthickness=0)
        detail_card.configure(highlightbackground="#cbd5e1")
        detail_card.pack(fill="both", expand=True)
        
        self.lbl_detail_meta = tk.Label(detail_card, text="请在左侧点击段落以查看详情", font=BOLD_FONT, fg="#475569", bg="#f1f5f9", anchor="w", padx=8, pady=4)
        self.lbl_detail_meta.pack(fill="x")
        
        # 1. 底部基础操作栏（打开文件、打开文件夹）
        btn_bar = tk.Frame(detail_card, bg="#ffffff", padx=8, pady=8)
        btn_bar.pack(fill="x", side="bottom")
        
        self.btn_open_file = tk.Button(
            btn_bar, 
            text="一键打开原文 🚀", 
            command=self._open_selected_file, 
            font=BOLD_FONT, 
            bg="#0f172a", 
            fg="#ffffff", 
            relief="flat",
            activebackground="#1e293b",
            cursor="hand2",
            padx=10,
            pady=4
        )
        self.btn_open_file.pack(side="left")
        self.btn_open_file.configure(state="disabled")
        
        self.btn_open_dir = tk.Button(
            btn_bar, 
            text="打开所在文件夹 📂", 
            command=self._open_selected_dir, 
            font=DEFAULT_FONT, 
            bg="#f1f5f9", 
            fg="#0f172a", 
            relief="flat",
            activebackground="#cbd5e1",
            cursor="hand2",
            padx=10,
            pady=4
        )
        self.btn_open_dir.pack(side="right")
        self.btn_open_dir.configure(state="disabled")
        
        # 2. 底部导出操作栏（导出当前段落、导出全部检索结果）
        export_bar = tk.Frame(detail_card, bg="#ffffff", padx=8, pady=4)
        export_bar.pack(fill="x", side="bottom", pady=(0, 8))
        
        self.btn_export_sel = tk.Button(
            export_bar, 
            text="导出当前段落 (Word) 📄", 
            command=self._export_selected_paragraph, 
            font=DEFAULT_FONT, 
            bg="#0284c7", 
            fg="#ffffff", 
            relief="flat",
            activebackground="#0369a1",
            activeforeground="#ffffff",
            cursor="hand2",
            padx=10,
            pady=4
        )
        self.btn_export_sel.pack(side="left")
        self.btn_export_sel.configure(state="disabled")
        
        self.btn_export_all = tk.Button(
            export_bar, 
            text="导出检索结果 (Word) 📂", 
            command=self._export_search_results, 
            font=DEFAULT_FONT, 
            bg="#10b981", 
            fg="#ffffff", 
            relief="flat",
            activebackground="#047857",
            activeforeground="#ffffff",
            cursor="hand2",
            padx=10,
            pady=4
        )
        self.btn_export_all.pack(side="right")
        self.btn_export_all.configure(state="disabled")
        
        # 3. 中间填充的文本正文（最后 pack，它将自适应填充上方与下方控件之间的全部空隙！）
        self.txt_detail_content = tk.Text(
            detail_card, 
            wrap="word", 
            font=DEFAULT_FONT, 
            bg="#ffffff", 
            fg="#334155", 
            bd=0, 
            highlightthickness=0, 
            padx=8, 
            pady=8
        )
        self.txt_detail_content.pack(fill="both", expand=True)
        self.txt_detail_content.insert("1.0", "暂无选中项。")
        self.txt_detail_content.configure(state="disabled")
        
        self.selected_item = None
        
        # 4. 最底部状态栏
        self.lbl_status = tk.Label(
            self.root, 
            text="就绪", 
            font=("Microsoft YaHei", 9), 
            fg="#64748b", 
            bg=self.bg_color, 
            anchor="w", 
            padx=15, 
            pady=4
        )
        self.lbl_status.pack(fill="x", side="bottom")

    # ==========================================
    # 5. UI 事件交互与线程控制
    # ==========================================

    def _update_db_stats(self):
        """刷新索引数据库的全局文件与段落总数"""
        try:
            fc, pc = self.db.get_stats()
            self.lbl_status.configure(text=f"本地索引状态：已成功索引文件数: {fc} 个，关联段落数: {pc} 个。")
        except Exception as e:
            self.lbl_status.configure(text=f"读取索引数据库失败: {e}")

    def _select_folder(self):
        """选择文件夹"""
        dir_path = filedialog.askdirectory()
        if dir_path:
            self.selected_folder = os.path.normpath(dir_path)
            self.dir_var.set(self.selected_folder)
            self._save_settings()  # 保存配置
            self._rebuild_category_ui()  # 重构子文件夹分类筛选列表

    def _start_indexing(self):
        """启动后台线程扫描文件夹并建立索引"""
        folder = self.dir_var.get()
        if not folder or folder == "请选择一个本地文件夹并建立索引..." or not os.path.exists(folder):
            messagebox.showwarning("提示", "请选择一个有效的本地文件夹！")
            return
            
        if self.is_indexing:
            messagebox.showinfo("正在扫描", "文件夹扫描建立索引任务正在运行中...")
            return
            
        # UI 变更为“正在索引”状态
        self.is_indexing = True
        self.progress_frame.pack(fill="x", padx=15, pady=(0, 8))
        self.btn_index.configure(state="disabled", text="正在扫描 ⌛")
        self.progress_bar["value"] = 0
        
        # 启动线程
        t = threading.Thread(target=self._index_folder_worker, args=(folder,), daemon=True)
        t.start()

    def _index_folder_worker(self, folder_path):
        """后台索引工作线程，提取文本并写入 SQLite"""
        try:
            self.indexing_queue.put(("status", "正在检索目录中的文档列表..."))
            
            # 支持的后缀
            allowed_exts = {'.docx', '.odt', '.pdf'}
            files_to_process = []
            
            # 递归遍历目录
            for root_dir, _, files in os.walk(folder_path):
                for f in files:
                    _, ext = os.path.splitext(f.lower())
                    if ext in allowed_exts:
                        filepath = os.path.join(root_dir, f)
                        files_to_process.append(filepath)
                        
            total_files = len(files_to_process)
            self.indexing_queue.put(("max_progress", total_files))
            
            if total_files == 0:
                self.indexing_queue.put(("finished", (0, 0)))
                return
                
            # 读取本地已索引的缓存
            db_cache = self.db.get_indexed_files()
            
            processed_count = 0
            new_or_updated = 0
            
            for fpath in files_to_process:
                processed_count += 1
                fname = os.path.basename(fpath)
                
                # 获取文件的当前修改时间
                try:
                    mtime = os.path.getmtime(fpath)
                except OSError:
                    continue
                    
                # 计算相对路径，仅将第一级子目录作为分类，其余深层嵌套的子目录合并到对应的一级子目录下
                norm_folder = os.path.abspath(folder_path).lower().replace("/", "\\")
                norm_file = os.path.abspath(fpath).lower().replace("/", "\\")
                if norm_file.startswith(norm_folder):
                    rel = fpath[len(folder_path):].strip("\\/")
                    # 按照系统路径分隔符拆分，取第一部分作为分类名称
                    parts = rel.split("\\") if "\\" in rel else rel.split("/")
                    category = parts[0] if len(parts) > 1 else "根目录"
                else:
                    category = "根目录"
                    
                # 检查修改时间及分类，判断是否需要增量重爬或直接更新分类
                if fpath in db_cache:
                    cached_mtime, cached_cat = db_cache[fpath]
                    if abs(cached_mtime - mtime) < 0.1:
                        if cached_cat == category:
                            self.indexing_queue.put(("progress", (processed_count, f"跳过未更改文件: {fname}")))
                            continue
                        else:
                            # 仅分类变化（如路径重命名或移动），快速更新分类信息
                            self.db.update_file_category(fpath, category)
                            self.indexing_queue.put(("progress", (processed_count, f"更新分类: {fname}")))
                            continue
                            
                # 需要提取文本并录入
                new_or_updated += 1
                self.indexing_queue.put(("progress", (processed_count, f"正在解析 [{new_or_updated}]: {fname}")))
                
                paragraphs = extract_text_from_file(fpath)
                if paragraphs:
                    self.db.add_file(fpath, paragraphs, category)
                    
            # 扫描并清理在物理硬盘上已被删除，但在数据库中还有残留的文件记录
            self.indexing_queue.put(("status", "正在比对并同步清理已被移除的文档记录..."))
            for cached_path in db_cache:
                if not os.path.exists(cached_path):
                    self.db.remove_file(cached_path)
                    
            self.indexing_queue.put(("finished", (total_files, new_or_updated)))
            
        except Exception as e:
            self.indexing_queue.put(("error", str(e)))

    def _listen_queue(self):
        """在主线程（UI线程）中以毫秒级轮询队列，以便更新 UI"""
        try:
            while True:
                msg_type, data = self.indexing_queue.get_nowait()
                if msg_type == "status":
                    self.lbl_progress_status.configure(text=data)
                elif msg_type == "max_progress":
                    self.progress_bar["maximum"] = data
                elif msg_type == "progress":
                    val, text = data
                    self.progress_bar["value"] = val
                    self.lbl_progress_status.configure(text=text)
                elif msg_type == "finished":
                    total, updated = data
                    self.is_indexing = False
                    self.progress_frame.pack_forget()
                    self.btn_index.configure(state="normal", text="更新/扫描索引 ⚡")
                    self._update_db_stats()
                    self._rebuild_category_ui()  # 扫描及入库完毕后，重构分类筛选列表
                    messagebox.showinfo("索引完成", f"同步索引完成！\n总文件: {total}，新导入/已更新: {updated}。")
                elif msg_type == "error":
                    self.is_indexing = False
                    self.progress_frame.pack_forget()
                    self.btn_index.configure(state="normal", text="更新/扫描索引 ⚡")
                    messagebox.showerror("索引出错", f"数据库增量更新过程中发生错误:\n{data}")
        except queue.Empty:
            pass
        finally:
            # 100ms 后再次调用，循环检查
            self.root.after(100, self._listen_queue)

    # ==========================================
    # 6. 搜索、评分与结果显示
    # ==========================================

    def _perform_search(self):
        """执行搜索并重绘左侧结果和右侧分析区"""
        term = self.search_var.get().strip()
        if not term:
            messagebox.showwarning("提示", "请输入要检索的术语！")
            return
        start_time = time.perf_counter()
        checked_cats = [cat for cat, var in self.category_vars.items() if var.get()]
        if self.category_vars and not checked_cats:
            self.results_frame.clear()
            self._clear_details()
            self.lbl_search_stats.configure(text="检索耗时: 0 ms | 找到 0 条结果")
            tk.Label(self.results_frame.scrollable_frame,
                     text="当前未勾选任何子文件夹分类。请先在右侧勾选需要检索的分类！",
                     font=DEFAULT_FONT, fg="#94a3b8", bg="#f8fafc", pady=100).pack(fill="x")
            return
        assoc2 = self.assoc2_var.get().strip() if hasattr(self, "assoc2_var") else ""
        assoc3 = self.assoc3_var.get().strip() if hasattr(self, "assoc3_var") else ""
        raw_results = self.db.search_paragraphs(term, checked_cats if self.category_vars else None)
        self.current_results = analyze_and_rank_results(raw_results, term, assoc2, assoc3)
        duration = (time.perf_counter() - start_time) * 1000
        count = len(self.current_results)
        self.lbl_search_stats.configure(text=f"检索耗时: {duration:.1f} ms | 找到 {count} 条结果")
        self.results_frame.clear()
        self._clear_details()
        self._displayed_count = 0
        if count == 0:
            tk.Label(self.results_frame.scrollable_frame,
                     text=f"未找到关于 '{term}' 的任何段落匹配，请确认已扫描并索引对应文件。",
                     font=DEFAULT_FONT, fg="#94a3b8", bg="#f8fafc", pady=100).pack(fill="x")
            self.btn_export_all.configure(state="disabled")
            return
        self.btn_export_all.configure(state="normal")
        self._load_more_results(page_size=50)
        if self.current_results:
            self._select_detail_item(self.current_results[0])

    def _load_more_results(self, page_size=50):
        """加载下一页 page_size 条结果到列表中"""
        if not hasattr(self, "_displayed_count"):
            self._displayed_count = 0
        term = self.search_var.get().strip()
        start = self._displayed_count
        end = min(len(self.current_results), start + page_size)
        batch = self.current_results[start:end]
        self._render_search_results_incremental(batch, start_idx=0, batch_size=10)
        self._displayed_count = end
        remaining = len(self.current_results) - self._displayed_count
        if remaining > 0:
            load_btn = tk.Button(
                self.results_frame.scrollable_frame,
                text=f"▼ 继续加载下 50 条（还剩 {remaining} 条）",
                font=DEFAULT_FONT, bg="#e0f2fe", fg="#0369a1",
                relief="flat", cursor="hand2", pady=6,
                command=lambda btn=None: self._on_load_more(load_btn)
            )
            load_btn.pack(fill="x", padx=10, pady=4)
        else:
            tk.Label(self.results_frame.scrollable_frame,
                     text=f"—— 已显示全部 {len(self.current_results)} 条结果 ——",
                     font=("Microsoft YaHei", 9, "italic"),
                     fg="#94a3b8", bg="#f8fafc", pady=10).pack(fill="x")

    def _on_load_more(self, old_btn):
        old_btn.destroy()
        self._load_more_results(page_size=50)


    def _render_search_results_incremental(self, display_results, start_idx=0, batch_size=10):
        """增量分批渲染检索结果卡片"""
        if not display_results:
            return
        end_idx = min(len(display_results), start_idx + batch_size)
        term = self.search_var.get().strip()
        for i in range(start_idx, end_idx):
            res = display_results[i]
            card = ModernCard(
                self.results_frame.scrollable_frame, res, term,
                self._select_detail_item, self._open_filepath_with_location
            )
            card.pack(fill="x", padx=10, pady=6)
        if end_idx < len(display_results):
            self.root.after(10, lambda: self._render_search_results_incremental(
                display_results, end_idx, batch_size))


    def _select_detail_item(self, result):
        """选择条目，填充到右侧侧边栏详情面板，并对判定句和术语进行醒目高亮"""
        self.selected_item = result
        term = self.search_var.get().strip()
        
        meta_text = f" 📄 {result['filename']} | 得分: {result['score']:.1f}"
        self.lbl_detail_meta.configure(text=meta_text)
        
        self.txt_detail_content.configure(state="normal")
        self.txt_detail_content.delete("1.0", "end")
        
        # 插入原文（始终显示整段原文，方便用户通读上下文）
        content = result["content"]
        self.txt_detail_content.insert("1.0", content + "\n\n")
        
        # 对段落文本中的检索术语及紧邻特征词进行精确高亮（黄底），关联词蓝底
        highlight_term_and_close_keywords(
            self.txt_detail_content,
            term,
            "detail_query",
            "detail_key_sent"
        )
        # 关联词高亮（橙底）
        assoc2 = self.assoc2_var.get().strip() if hasattr(self, 'assoc2_var') else ""
        assoc3 = self.assoc3_var.get().strip() if hasattr(self, 'assoc3_var') else ""
        for assoc_term in [assoc2, assoc3]:
            if assoc_term:
                highlight_term_in_widget(self.txt_detail_content, assoc_term, "assoc_hl")
        
        # 添加评分说明
        self.txt_detail_content.insert("end", "────────────────────────────────\n", "divider")
        self.txt_detail_content.insert("end", f"📄 原文路径: {result['filepath']}\n", "meta")
        
        # 提取书名和篇目信息，格式化展示，不显示第几段
        filename = result['filename']
        book_match = re.search(r'《([^》]+)》', filename)
        book_title = book_match.group(1) if book_match else os.path.splitext(filename)[0]
        
        # 仅显示二级标题篇目，不显示大标题或小标题
        clean_heading = result['heading'].split(" > ")[-1].strip()
        if clean_heading and clean_heading != "（无标题）":
            title_info = f"《{book_title}·{clean_heading}》"
        else:
            title_info = f"《{book_title}》"
            
        self.txt_detail_content.insert("end", f"📖 所在篇目: {title_info}\n", "meta")
        
        if result["info"]:
            self.txt_detail_content.insert("end", f"⭐ 评分特征: {' | '.join(result['info'])}\n", "meta")
            self.txt_detail_content.insert("end", f"💡 判定句: \"{result['matched_sentence']}\"", "highlight_sent")
            
        self.txt_detail_content.tag_configure("divider", foreground="#e2e8f0")
        self.txt_detail_content.tag_configure("meta", font=("Microsoft YaHei", 9), foreground="#64748b")
        self.txt_detail_content.tag_configure("highlight_sent", font=("Microsoft YaHei", 9, "italic"), foreground="#0ea5e9")
        self.txt_detail_content.tag_configure("assoc_hl", background="#4ade80")  # 鲜绿色关联词底色
        
        self.txt_detail_content.configure(state="disabled")
        
        # 启用操作按钮
        self.btn_open_file.configure(state="normal")
        self.btn_open_dir.configure(state="normal")
        self.btn_export_sel.configure(state="normal")

    def _clear_details(self):
        """重置右侧详情面板"""
        self.selected_item = None
        self.lbl_detail_meta.configure(text="请在左侧点击段落以查看详情")
        self.txt_detail_content.configure(state="normal")
        self.txt_detail_content.delete("1.0", "end")
        self.txt_detail_content.insert("1.0", "暂无选中项。")
        self.txt_detail_content.configure(state="disabled")
        self.btn_open_file.configure(state="disabled")
        self.btn_open_dir.configure(state="disabled")
        self.btn_export_sel.configure(state="disabled")

    # ==========================================
    # 7. 一键调用系统默认程序打开原文
    # ==========================================

    def _open_selected_file(self):
        """调用系统程序打开当前选中的文档并定位到对应段落"""
        if self.selected_item:
            self._open_filepath_with_location(self.selected_item["filepath"], self.selected_item["content"])

    def _open_selected_dir(self):
        """打开文件所在的目录"""
        if self.selected_item:
            dir_path = os.path.dirname(self.selected_item["filepath"])
            self._open_filepath_with_location(dir_path)

    def _open_filepath_with_location(self, path, paragraph_text=None):
        """安全的跨平台/Windows 打开系统程序逻辑，支持 Word 定位段落"""
        path = os.path.normpath(path)
        if not os.path.exists(path):
            messagebox.showerror("打开失败", f"文件或文件夹已不存在！\n路径: {path}")
            return
            
        _, ext = os.path.splitext(path.lower())
        if ext == '.docx' and paragraph_text and HAS_WIN32COM:
            # 开启后台线程进行 COM 定位，防止拉起 Word 导致 Tkinter 界面出现短时假死
            t = threading.Thread(target=self._locate_docx_thread, args=(path, paragraph_text), daemon=True)
            t.start()
        else:
            try:
                os.startfile(path)
            except Exception as e:
                messagebox.showerror("打开出错", f"无法调用系统默认程序打开: {e}")

    def _locate_docx_thread(self, filepath, paragraph_text):
        """在后台线程中使用 COM 查找并选中目标文本以实现自动定位"""
        success = locate_and_select_in_word(filepath, paragraph_text)
        if not success:
            # 定位失败兜底：使用系统关联的默认程序正常打开文档
            try:
                os.startfile(filepath)
            except Exception:
                pass

    # ==========================================
    # 8. 子文件夹分类筛选交互逻辑与持久化
    # ==========================================

    def _load_settings(self):
        """加载已保存的历史选择文件夹配置"""
        self.selected_folder = ""
        settings_path = os.path.join(EXE_DIR, "settings.json")
        if os.path.exists(settings_path):
            try:
                with open(settings_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    self.selected_folder = data.get("selected_folder", "")
            except Exception as e:
                print(f"加载持久化配置失败: {e}")

    def _save_settings(self):
        """保存当前选中的文件夹配置"""
        settings_path = os.path.join(EXE_DIR, "settings.json")
        try:
            with open(settings_path, "w", encoding="utf-8") as f:
                json.dump({"selected_folder": self.selected_folder}, f, ensure_ascii=False, indent=4)
        except Exception as e:
            print(f"保存持久化配置失败: {e}")

    def _rebuild_category_ui(self):
        """根据当前数据库中的文件分类，重新生成右侧的分类勾选复选框"""
        self.category_frame.clear()
        self.category_vars = {}
        
        categories = self.db.get_all_categories()
        if not categories:
            lbl_none = tk.Label(self.category_frame.scrollable_frame, text="暂无分类数据，请更新索引。", font=DEFAULT_FONT, fg="#94a3b8", bg="#f8fafc")
            lbl_none.pack(padx=10, pady=20)
            return

        # 判断是否存在多个分类（即除了根目录外还有其他子文件夹）
        has_multiple_categories = len(categories) > 1

        # 逐个创建各分类的 Checkbutton 变量
        for cat in categories:
            # 当存在多个子文件夹时，根目录默认不勾选 (False)；其余子文件夹默认勾选 (True)
            if cat == "根目录" and has_multiple_categories:
                default_val = False
            else:
                default_val = True
            self.category_vars[cat] = tk.BooleanVar(value=default_val)

        # 计算全选 (All) 的状态
        all_checked = all(var.get() for var in self.category_vars.values())
        self.var_select_all = tk.BooleanVar(value=all_checked)

        cb_all = tk.Checkbutton(
            self.category_frame.scrollable_frame, 
            text="全选 (All)", 
            variable=self.var_select_all,
            command=self._on_select_all_clicked,
            font=BOLD_FONT,
            bg="#f8fafc",
            fg="#0f172a",
            activebackground="#f8fafc",
            anchor="w"
        )
        cb_all.pack(fill="x", padx=10, pady=2)

        for cat in categories:
            var = self.category_vars[cat]
            cb = tk.Checkbutton(
                self.category_frame.scrollable_frame, 
                text=cat, 
                variable=var,
                command=self._on_category_cb_clicked,
                font=DEFAULT_FONT,
                bg="#f8fafc",
                fg="#334155",
                activebackground="#f8fafc",
                anchor="w"
            )
            cb.pack(fill="x", padx=20, pady=1)

    def _on_select_all_clicked(self):
        val = self.var_select_all.get()
        for var in self.category_vars.values():
            var.set(val)
        # 自动触发重新搜索（如果搜索框有值）
        if self.search_var.get().strip():
            self._perform_search()
            
    def _on_category_cb_clicked(self):
        # 如果所有子分类都是 True，则全选是 True，否则是 False
        all_checked = all(var.get() for var in self.category_vars.values())
        self.var_select_all.set(all_checked)
        # 自动触发重新搜索
        if self.search_var.get().strip():
            self._perform_search()

    # ==========================================
    # 9. 格式化段落与搜索结果导出 Word 文档
    # ==========================================

    def _export_selected_paragraph(self):
        """将当前选中的单个段落（保持关键句红色加粗与检索词黄底格式）导出至 Word"""
        if not self.selected_item:
            messagebox.showwarning("提示", "当前未选中任何段落！")
            return
        
        if docx is None:
            messagebox.showerror("导出失败", "系统检测到未安装 python-docx，无法导出 Word！请运行 pip install python-docx")
            return
            
        file_path = filedialog.asksaveasfilename(
            title="导出段落到 Word",
            defaultextension=".docx",
            filetypes=[("Word 文档", "*.docx")],
            initialfile=f"导出段落_{self.selected_item['filename']}.docx"
        )
        if not file_path:
            return
            
        try:
            doc = docx.Document()
            # 设置左右页边距为 2cm
            for section in doc.sections:
                section.left_margin = Cm(2.0)
                section.right_margin = Cm(2.0)
                
            # 设置 Normal 样式的默认字体为 宋体，字号为 小四 (12pt)
            normal_style = doc.styles['Normal']
            normal_style.font.name = '宋体'
            normal_style.font.size = Pt(12)
            rPr = normal_style.element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 添加大纲标题
            h = doc.add_heading("本地 Word 术语检索段落导出", level=1)
            # 标题也统一应用宋体
            for run in h.runs:
                run.font.name = '宋体'
                run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')
            
            p_file = doc.add_paragraph(f"📄 来源文档: {self.selected_item['filepath']}")
            p_pos = doc.add_paragraph(f"📌 段落位置: 文件第 {self.selected_item['para_index'] + 1} 个自然段 | 标题: {self.selected_item['heading']}")
            p_info = doc.add_paragraph(f"⭐ 评分特征: {' | '.join(self.selected_item['info'])}")
            doc.add_paragraph("────────────────────────────────")
            
            # 文本正文段落，设置首行缩进两字符 (小四号字宽为 12pt，两字符为 24pt)
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(24)
            
            query = self.search_var.get().strip()
            assoc2 = self.assoc2_var.get().strip() if hasattr(self, 'assoc2_var') else ""
            assoc3 = self.assoc3_var.get().strip() if hasattr(self, 'assoc3_var') else ""

            src_p = None
            if self.selected_item['filepath'].lower().endswith('.docx'):
                src_p = get_source_docx_paragraph(self.selected_item['filepath'], self.selected_item['para_index'])

            if src_p:
                add_formatted_text_from_src_paragraph(
                    p, src_p, query,
                    self.selected_item.get("has_close_definition", False),
                    assoc2=assoc2, assoc3=assoc3
                )
            else:
                add_formatted_text_to_docx_paragraph(
                    p, self.selected_item["content"], query,
                    has_close_def=self.selected_item.get("has_close_definition", False),
                    assoc2=assoc2, assoc3=assoc3
                )
            
            doc.save(file_path)
            # 自动打开生成的文档
            try:
                os.startfile(file_path)
            except Exception:
                pass
            messagebox.showinfo("成功", f"段落已成功导出并已自动打开：\n{file_path}")
        except Exception as e:
            messagebox.showerror("导出失败", f"导出 Word 过程中发生错误:\n{e}")

    def _export_search_results(self):
        """将当前检索到的结果列表（前 50 条）批量导出至单个 Word 文档（格式高亮完整保留）"""
        if not self.current_results:
            messagebox.showwarning("提示", "当前检索结果为空！")
            return
            
        if docx is None:
            messagebox.showerror("导出失败", "系统检测到未安装 python-docx，无法导出 Word！请运行 pip install python-docx")
            return
        
        query = self.search_var.get().strip()
        file_path = filedialog.asksaveasfilename(
            title="导出检索结果到 Word",
            defaultextension=".docx",
            filetypes=[("Word 文档", "*.docx")],
            initialfile=f"检索结果_{query}.docx"
        )
        if not file_path:
            return
            
        try:
            doc = docx.Document()
            # 设置左右页边距为 2cm
            for section in doc.sections:
                section.left_margin = Cm(2.0)
                section.right_margin = Cm(2.0)
                
            # 设置 Normal 样式的默认字体为 宋体，字号为 小四 (12pt)
            normal_style = doc.styles['Normal']
            normal_style.font.name = '宋体'
            normal_style.font.size = Pt(12)
            rPr = normal_style.element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), '宋体')
            
            h = doc.add_heading(f"术语 '{query}' 检索结果导出", level=1)
            for run in h.runs:
                run.font.name = '宋体'
                run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')
                
            count = len(self.current_results)
            # 超过300条固定导出前300条高分结果，100条以内全部导出，中间段询问
            if count <= 100:
                export_limit = count
            elif count <= 300:
                answer = messagebox.askyesno(
                    "导出数量",
                    f"共找到 {count} 条结果。\n"
                    f"· 点击《是》：导出全部 {count} 条\n"
                    f"· 点击《否》：仅导出评分最高的 100 条"
                )
                export_limit = count if answer else 100
            else:
                export_limit = 300
                messagebox.showinfo("导出数量", f"共找到 {count} 条结果，将导出评分最高的前 300 条。")
            
            doc.add_paragraph(f"🔍 全文匹配共计 {count} 条，以下根据相关度评分排序导出前 {export_limit} 条：")
            doc.add_paragraph("────────────────────────────────")
            
            assoc2 = self.assoc2_var.get().strip() if hasattr(self, 'assoc2_var') else ""
            assoc3 = self.assoc3_var.get().strip() if hasattr(self, 'assoc3_var') else ""

            # 缓存已打开的 docx 文档对象及预查好的段落列表，避免重复打开与遍历（极速导出）
            # 结构: docx_cache[fpath] = (d_obj, p_objs_list)
            docx_cache = {}

            for i, item in enumerate(self.current_results[:export_limit]):
                p_meta = doc.add_paragraph()
                run_num = p_meta.add_run(f"[{i+1}] ")
                run_num.bold = True
                run_num.font.name = '宋体'
                run_num._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')

                filename = item['filename']
                book_match = re.search(r'《([^》]+)》', filename)
                book_title = book_match.group(1) if book_match else os.path.splitext(filename)[0]
                clean_heading = item['heading'].split(" > ")[-1].strip()
                if clean_heading and clean_heading != "（无标题）":
                    title_info = f"《{book_title}·{clean_heading}》"
                else:
                    title_info = f"《{book_title}》"

                run_desc = p_meta.add_run(f"📄 {title_info} | 得分: {item['score']:.1f}\n")
                run_desc.italic = True
                run_desc.font.name = '宋体'
                run_desc._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')

                p_content = doc.add_paragraph()
                p_content.paragraph_format.first_line_indent = Pt(24)

                src_p = None
                fpath = item['filepath']
                if fpath.lower().endswith('.docx'):
                    if fpath not in docx_cache:
                        try:
                            d_obj = docx.Document(fpath)
                            p_objs = [p for p in d_obj.paragraphs if p.text.strip()]
                            for tbl in d_obj.tables:
                                for row in tbl.rows:
                                    for cell in row.cells:
                                        for p in cell.paragraphs:
                                            if p.text.strip():
                                                p_objs.append(p)
                            docx_cache[fpath] = (d_obj, p_objs)
                        except Exception:
                            docx_cache[fpath] = (None, [])
                    _, p_objs = docx_cache[fpath]
                    p_idx = item['para_index']
                    if 0 <= p_idx < len(p_objs):
                        src_p = p_objs[p_idx]

                if src_p:
                    add_formatted_text_from_src_paragraph(
                        p_content, src_p, query,
                        item.get("has_close_definition", False),
                        assoc2=assoc2, assoc3=assoc3
                    )
                else:
                    add_formatted_text_to_docx_paragraph(
                        p_content, item["content"], query,
                        has_close_def=item.get("has_close_definition", False),
                        assoc2=assoc2, assoc3=assoc3
                    )

                doc.add_paragraph("─" * 50)

            doc.save(file_path)
            # 自动打开生成的文档
            try:
                if sys.platform == "win32":
                    os.startfile(file_path)
                elif sys.platform == "darwin": # Mac
                    import subprocess
                    subprocess.run(["open", file_path])
                else:
                    import subprocess
                    subprocess.run(["xdg-open", file_path])
            except Exception:
                pass
            messagebox.showinfo("成功", f"检索结果已成功导出并已自动打开：\n{file_path}")
        except Exception as e:
            messagebox.showerror("导出失败", f"导出 Word 过程中发生错误:\n{e}")

    def _batch_term_search_and_export(self):
        """导入一个包含多个术语的文件（每行一个），批量检索并在保留原格式的基础上导出为 Word 报告"""
        if docx is None:
            messagebox.showerror("提示", "系统检测到未安装 python-docx，无法使用此功能！请运行 pip install python-docx")
            return
            
        # 1. 选择包含术语的文档
        term_file = filedialog.askopenfilename(
            title="选择包含术语的文档 (TXT 或 DOCX)",
            filetypes=[("支持的文档", "*.txt;*.docx"), ("文本文件", "*.txt"), ("Word 文档", "*.docx")]
        )
        if not term_file:
            return
            
        # 读取所有术语
        terms = []
        _, ext = os.path.splitext(term_file.lower())
        try:
            if ext == '.txt':
                with open(term_file, 'r', encoding='utf-8', errors='ignore') as f:
                    for line in f:
                        t = line.strip()
                        if t:
                            terms.append(t)
            elif ext == '.docx':
                doc = docx.Document(term_file)
                for p in doc.paragraphs:
                    t = p.text.strip()
                    if t:
                        terms.append(t)
        except Exception as e:
            messagebox.showerror("读取失败", f"无法读取术语文件内容:\n{e}")
            return
            
        if not terms:
            messagebox.showwarning("提示", "未在该文档中找到任何术语（每行一个）！")
            return
            
        # 2. 选择导出报告的保存路径
        save_path = filedialog.asksaveasfilename(
            title="选择保存批量检索报告的位置",
            defaultextension=".docx",
            filetypes=[("Word 文档", "*.docx")],
            initialfile="批量术语检索报告.docx"
        )
        if not save_path:
            return
            
        # 3. 开始批量检索和写入报告
        try:
            doc = docx.Document()
            # 设置左右页边距为 2cm
            for section in doc.sections:
                section.left_margin = Cm(2.0)
                section.right_margin = Cm(2.0)
                
            # 设置 Normal 样式的默认字体为 宋体，字号为 小四 (12pt)
            normal_style = doc.styles['Normal']
            normal_style.font.name = '宋体'
            normal_style.font.size = Pt(12)
            rPr = normal_style.element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 主标题
            h_main = doc.add_heading("批量术语检索报告", level=1)
            for run in h_main.runs:
                run.font.name = '宋体'
                run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')
                
            doc.add_paragraph(f"📂 术语来源文档: {os.path.basename(term_file)}")
            doc.add_paragraph(f"📊 共计导入术语 {len(terms)} 个，针对每个术语提取匹配度最高（定义优先）的段落（若搜索结果超过 400 处提取前 200 处，否则提取前 100 处）：")
            doc.add_paragraph("────────────────────────────────")
            
            processed_count = 0
            
            for t_idx, term in enumerate(terms):
                # 检索数据库
                raw_results = self.db.search_paragraphs(term)
                if not raw_results:
                    h_term = doc.add_heading(f"[{t_idx+1}] 术语: {term} (未匹配到结果)", level=2)
                    for run in h_term.runs:
                        run.font.name = '宋体'
                        run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')
                    doc.add_paragraph("在当前索引资料库中未找到任何包含此术语的段落。")
                    doc.add_paragraph("─" * 40)
                    continue
                    
                # 评分与排序
                ranked_results = analyze_and_rank_results(raw_results, term)
                
                h_term = doc.add_heading(f"[{t_idx+1}] 术语: {term} (匹配 {len(ranked_results)} 条)", level=2)
                for run in h_term.runs:
                    run.font.name = '宋体'
                    run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')
                
                match_count = len(ranked_results)
                # 计算导出数量限制：结果数 >= 400 时导出 200 条，否则至少 100 条（或全部）
                export_limit = 200 if match_count >= 400 else min(match_count, 100)
                
                # 写入前 export_limit 条结果（提取深度扩充）
                for r_idx, item in enumerate(ranked_results[:export_limit]):
                    p_meta = doc.add_paragraph()
                    run_num = p_meta.add_run(f"  ({r_idx+1}) ")
                    run_num.bold = True
                    run_num.font.name = '宋体'
                    run_num._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')
                    
                    filename = item['filename']
                    book_match = re.search(r'《([^》]+)》', filename)
                    book_title = book_match.group(1) if book_match else os.path.splitext(filename)[0]
                    clean_heading = item['heading'].split(" > ")[-1].strip()
                    if clean_heading and clean_heading != "（无标题）":
                        title_info = f"《{book_title}·{clean_heading}》"
                    else:
                        title_info = f"《{book_title}》"
                    
                    run_desc = p_meta.add_run(f"📄 {title_info} | 得分: {item['score']:.1f}\n")
                    run_desc.italic = True
                    run_desc.font.name = '宋体'
                    run_desc._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')
                    
                    # 正文段落，设置首行缩进两字符 (小四号字宽为 12pt，两字符为 24pt)
                    p_content = doc.add_paragraph()
                    p_content.paragraph_format.first_line_indent = Pt(24)
                    
                    # 尝试读取源 DOCX 段落以保留原始格式
                    src_p = None
                    if item['filepath'].lower().endswith('.docx'):
                        src_p = get_source_docx_paragraph(item['filepath'], item['para_index'])
                        
                    if src_p:
                        add_formatted_text_from_src_paragraph(
                            p_content, 
                            src_p, 
                            term, 
                            item.get("has_close_definition", False)
                        )
                    else:
                        add_formatted_text_to_docx_paragraph(
                            p_content, 
                            item["content"], 
                            term, 
                            has_close_def=item.get("has_close_definition", False)
                        )
                        
                # 添加节分割线
                doc.add_paragraph("─" * 50)
                processed_count += 1
                
            doc.save(save_path)
            messagebox.showinfo("批量导出成功", f"批量术语检索已完成！\n共处理了 {processed_count}/{len(terms)} 个有匹配项的术语。\n报告保存路径：\n{save_path}")
        except Exception as e:
            messagebox.showerror("批量处理失败", f"批量检索与生成 Word 过程中发生错误:\n{e}")

# ==========================================
# 10. Word 导出文字高亮生成辅助函数
# ==========================================

def get_filtered_paragraph_slices(text, query):
    """
    当检索词出现次数少于 4 次时，提取关键词前后各 30 个字符的范围。
    合并重叠区间后，用 ' ... ' 拼接。
    返回 (new_text, index_mapping)
    其中 index_mapping 是与 new_text 长度相同的列表，存放新字符在原 text 中的索引（若为合成字符如 ' ... ' 则为 None）
    """
    if not text or not query:
        return text, list(range(len(text)))
        
    n = len(text)
    spans = []
    q_len = len(query)
    start = 0
    query_lower = query.lower()
    text_lower = text.lower()
    
    while True:
        idx = text_lower.find(query_lower, start)
        if idx == -1:
            break
        spans.append((idx, idx + q_len))
        start = idx + q_len
        
    if not spans:
        return text, list(range(len(text)))
        
    # 构建每个匹配前后 30 字符的区间
    intervals = []
    for s_start, s_end in spans:
        left = max(0, s_start - 30)
        right = min(n, s_end + 30)
        intervals.append([left, right])
        
    # 合并区间
    intervals.sort(key=lambda x: x[0])
    merged = []
    for interval in intervals:
        if not merged:
            merged.append(interval)
        else:
            prev = merged[-1]
            if interval[0] <= prev[1] + 5:  # 允许 5 个字符内的微小间距合并，避免过于碎片化
                prev[1] = max(prev[1], interval[1])
            else:
                merged.append(interval)
                
    # 构造新文本和映射
    new_text_chars = []
    mapping = []
    
    # 如果第一个区间不是文档开头，先加前置省略号
    if merged[0][0] > 0:
        for c in "... ":
            new_text_chars.append(c)
            mapping.append(None)
            
    for idx, (m_start, m_end) in enumerate(merged):
        for i in range(m_start, m_end):
            new_text_chars.append(text[i])
            mapping.append(i)
        if idx < len(merged) - 1:
            for c in " ... ":
                new_text_chars.append(c)
                mapping.append(None)
                
    # 如果最后一个区间不是文档末尾，加后置省略号
    if merged[-1][1] < n:
        for c in " ...":
            new_text_chars.append(c)
            mapping.append(None)
            
    return "".join(new_text_chars), mapping

def highlight_term_in_widget(text_widget, term, tag_name):
    """在 Tkinter Text 控件中高亮指定单个术语（用于关联词橙底高亮）"""
    if not term:
        return
    start = "1.0"
    while True:
        pos = text_widget.search(term, start, stopindex="end")
        if not pos:
            break
        end_pos = f"{pos}+{len(term)}c"
        text_widget.tag_add(tag_name, pos, end_pos)
        start = end_pos


def add_formatted_text_to_docx_paragraph(docx_para, text, query, key_sentence=None, has_close_def=False, assoc2="", assoc3=""):
    """将包含检索术语和紧邻定义高亮格式的文本添加到 python-docx 段落中，并应用宋体、小四号字"""
    if not text:
        return
        
    n = len(text)
    char_attrs = []
    # 为每个字符分配属性: [is_red, is_yellow, is_orange]
    for _ in range(n):
        char_attrs.append([False, False, False])

    # 拆分同义词列表
    terms = [t.strip() for t in re.split(r'[|/,\s，]+', query) if t.strip()] or [query]

    # 1. 标记所有 terms 范围为 yellow (黄底)
    query_spans = []
    for q in terms:
        if not q:
            continue
        start = 0
        while True:
            idx = text.find(q, start)
            if idx == -1:
                break
            query_spans.append((idx, idx + len(q)))
            for i in range(idx, idx + len(q)):
                char_attrs[i][1] = True
            start = idx + len(q)

    # 2. 标记关联词为 orange (橙底)
    for assoc_q in [assoc2, assoc3]:
        if not assoc_q:
            continue
        start = 0
        while True:
            idx = text.find(assoc_q, start)
            if idx == -1:
                break
            for i in range(idx, idx + len(assoc_q)):
                char_attrs[i][2] = True
            start = idx + len(assoc_q)

    # 3. 标记紧邻的 definition keywords 为 red (仅当 has_close_def 为 True)
    if has_close_def and query_spans:
        type_a_kws = ["是指", "就是", "意味着", "关键是", "最重要", "是", "定义为", "定性为", "如何定义", "如何定性", "怎么定义", "怎么定性"]
        type_b_kws = ["什么是", "所谓", "属于", "就是", "如何定义", "如何定性", "怎么定义", "怎么定性", "定义", "定性"]
        
        for q_start, q_end in query_spans:
            # 2a. 优先检查组合模式：“所谓 [术语] 就是” (两者均紧挨着，距离均为 0)
            has_suowei_prefix = (q_start - 2 >= 0 and text[q_start - 2:q_start] == "所谓")
            has_jiushi_suffix = (q_end + 2 <= n and text[q_end:q_end + 2] == "就是")
            
            if has_suowei_prefix and has_jiushi_suffix:
                for i in range(q_start - 2, q_start):
                    char_attrs[i][0] = True
                for i in range(q_end, q_end + 2):
                    char_attrs[i][0] = True
                continue
                
            # 2b. 检查 Type A (紧邻，右侧距离为 0)
            for kw in type_a_kws:
                kw_len = len(kw)
                if q_end + kw_len <= n:
                    if text[q_end:q_end + kw_len] == kw:
                        for i in range(q_end, q_end + kw_len):
                            char_attrs[i][0] = True
                        break
            # 2c. 检查 Type B (紧邻，左侧距离为 0)
            for kw in type_b_kws:
                kw_len = len(kw)
                check_start = q_start - kw_len
                if check_start >= 0:
                    if text[check_start:check_start + kw_len] == kw:
                        for i in range(check_start, check_start + kw_len):
                            char_attrs[i][0] = True
                        break
                            
    # 合并连续相同属性字符块，添加为 Word run
    i = 0
    while i < n:
        is_red, is_yellow, is_orange = char_attrs[i]
        j = i + 1
        while j < n and char_attrs[j] == [is_red, is_yellow, is_orange]:
            j += 1
        run_text = text[i:j]
        run = docx_para.add_run(run_text)
        run.font.name = '宋体'
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
        if is_red:
            run.font.color.rgb = RGBColor(220, 38, 38)
            run.bold = True
        if is_yellow:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            run.bold = True
        if is_orange and not is_yellow:  # 鲜绿色关联词底色
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        i = j



def get_paragraph_from_cached_doc(cached_doc, para_index):
    """从已缓存的 docx.Document 对象中提取指定索引的段落，避免重复打开文件"""
    try:
        p_objects = [p for p in cached_doc.paragraphs if p.text.strip()]
        for table in cached_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            p_objects.append(p)
        if 0 <= para_index < len(p_objects):
            return p_objects[para_index]
    except Exception as e:
        print(f"从缓存文档提取段落失败: {e}")
    return None


def get_source_docx_paragraph(filepath, para_index):
    """尝试从源 .docx 文档中提取指定 index 的原始段落对象，用于后续保留源文件格式"""
    if docx is None or not os.path.exists(filepath):
        return None
    try:
        doc = docx.Document(filepath)
        p_objects = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                p_objects.append(p)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        text = p.text.strip()
                        if text:
                            p_objects.append(p)
        if 0 <= para_index < len(p_objects):
            return p_objects[para_index]
    except Exception as e:
        print(f"提取源 Word 段落对象失败 {filepath}: {e}")
    return None

def add_formatted_text_from_src_paragraph(dest_p, src_p, query, has_close_def, assoc2="", assoc3=""):
    """从源 paragraph (src_p) 复制文本及各种 XML 级格式（字体颜色、背景高亮/阴影等），并叠加高亮黄底与定义词红字"""
    if not src_p:
        return
        
    text = src_p.text
    if not text:
        return
        
    n = len(text)

    char_runs = []
    for run in src_p.runs:
        for _ in run.text:
            char_runs.append(run)

    if len(char_runs) != n:
        add_formatted_text_to_docx_paragraph(dest_p, text, query, has_close_def=has_close_def, assoc2=assoc2, assoc3=assoc3)
        return

    # 为每个字符分配属性: [is_red, is_yellow, is_orange]
    char_attrs = [[False, False, False] for _ in range(n)]

    # 拆分同义词
    terms = [t.strip() for t in re.split(r'[|/,\s，]+', query) if t.strip()] or [query]

    # 1. 标记所有 terms 为 yellow
    query_spans = []
    for q in terms:
        if not q:
            continue
        start = 0
        while True:
            idx = text.find(q, start)
            if idx == -1:
                break
            query_spans.append((idx, idx + len(q)))
            for i in range(idx, idx + len(q)):
                char_attrs[i][1] = True
            start = idx + len(q)

    # 2. 标记关联词为 orange
    for assoc_q in [assoc2, assoc3]:
        if not assoc_q:
            continue
        start = 0
        while True:
            idx = text.find(assoc_q, start)
            if idx == -1:
                break
            for i in range(idx, idx + len(assoc_q)):
                char_attrs[i][2] = True
            start = idx + len(assoc_q)

    # 3. 标记紧邻定义词为 red
    if has_close_def and query_spans:
        type_a_kws = ["是指", "就是", "意味着", "关键是", "最重要", "是", "定义为", "定性为", "如何定义", "如何定性", "怎么定义", "怎么定性"]
        type_b_kws = ["什么是", "所谓", "属于", "就是", "如何定义", "如何定性", "怎么定义", "怎么定性", "定义", "定性"]
        
        for q_start, q_end in query_spans:
            # 优先检查组合模式：“所谓 [术语] 就是” (两者均紧挨着，距离均为 0)
            has_suowei_prefix = (q_start - 2 >= 0 and text[q_start - 2:q_start] == "所谓")
            has_jiushi_suffix = (q_end + 2 <= n and text[q_end:q_end + 2] == "就是")
            
            if has_suowei_prefix and has_jiushi_suffix:
                for i in range(q_start - 2, q_start):
                    char_attrs[i][0] = True
                for i in range(q_end, q_end + 2):
                    char_attrs[i][0] = True
                continue
                
            # Type A (紧邻，右侧距离为 0)
            for kw in type_a_kws:
                kw_len = len(kw)
                if q_end + kw_len <= n:
                    if text[q_end:q_end + kw_len] == kw:
                        for i in range(q_end, q_end + kw_len):
                            char_attrs[i][0] = True
                        break
            # Type B (紧邻，左侧距离为 0)
            for kw in type_b_kws:
                kw_len = len(kw)
                check_start = q_start - kw_len
                if check_start >= 0:
                    if text[check_start:check_start + kw_len] == kw:
                        for i in range(check_start, check_start + kw_len):
                            char_attrs[i][0] = True
                        break
                        
    # 3. 合并相同 src_run 以及属性的相邻字符块并添加为 Run
    import copy
    i = 0
    while i < n:
        src_run = char_runs[i]
        is_red, is_yellow, is_orange = char_attrs[i]
        j = i + 1
        while j < n and char_runs[j] == src_run and char_attrs[j] == [is_red, is_yellow, is_orange]:
            j += 1
        run_text = text[i:j]
        if src_run:
            new_r = parse_xml(src_run._element.xml)
            dest_p._element.append(new_r)
            run = docx.text.run.Run(new_r, dest_p)
            run.text = run_text
        else:
            run = dest_p.add_run(run_text)
            run.font.name = '宋体'
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)
        if is_red:
            run.font.color.rgb = docx.shared.RGBColor(220, 38, 38)
            run.bold = True
        if is_yellow:
            run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
            run.bold = True
        if is_orange and not is_yellow:
            run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.BRIGHT_GREEN
        i = j

def locate_and_select_in_word(filepath, paragraph_text):
    """在 Word 中打开文件并搜索/选中指定段落文本以实现自动定位"""
    if not HAS_WIN32COM or not paragraph_text:
        return False
    try:
        # COM 初始化以防多线程报错
        pythoncom.CoInitialize()
        
        # 提取前 40 个字符作为查找条件，过滤特殊标点符号以提升匹配成功率，且不能超过 Word Find 255 字节限制
        search_text = paragraph_text.strip()
        search_text = re.sub(r'[\r\n\t]+', ' ', search_text)
        search_text = search_text[:40]
        if not search_text:
            return False
            
        # 尝试获取已有实例，没有则创建新实例
        try:
            word = win32com.client.GetActiveObject("Word.Application")
        except Exception:
            word = win32com.client.Dispatch("Word.Application")
            
        word.Visible = True
        doc = word.Documents.Open(os.path.abspath(filepath))
        
        # 将光标定位至文档开头
        word.Selection.Start = 0
        word.Selection.End = 0
        
        # 调用 Word 的 Find 查找逻辑
        find = word.Selection.Find
        find.ClearFormatting()
        find.Text = search_text
        find.Forward = True
        find.Wrap = 1  # wdFindAsk
        
        if find.Execute():
            # 查找到后，Word 会自动将光标移至该位置并选中。我们激活 Word 窗口让其置顶
            word.Activate()
            return True
    except Exception as e:
        print(f"Word COM 文本定位失败: {e}")
    return False

# ==========================================
# 8. 启动引导
# ==========================================

if __name__ == "__main__":
    # Windows 高 DPI 适配（防止字体模糊）
    try:
        from ctypes import windll
        windll.shcore.SetProcessDpiAwareness(1)
    except Exception:
        pass

    # 检查是否缺失必要依赖
    missing = []
    if docx is None:
        missing.append("python-docx (解析.docx Word文件)")
    if PdfReader is None:
        missing.append("pypdf (解析.pdf 文件)")
        
    root = tk.Tk()
    
    # 警告弹出
    if missing:
        msg = "部分文件解析库未安装，对应格式将无法提取文本：\n\n" + \
              "\n".join([f"• {m}" for m in missing]) + \
              "\n\n建议您根据需要使用以下命令安装：\n" + \
              "pip install python-docx pypdf"
        messagebox.showinfo("解析库提示", msg)
        
    app = TermAnalyzerApp(root)
    
    # 运行主循环，并在关闭时清理连接
    def on_closing():
        app.db.close()
        root.destroy()
        
    root.protocol("WM_DELETE_WINDOW", on_closing)
    root.mainloop()
